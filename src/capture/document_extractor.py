"""
Async Document Extraction Module

Provides efficient async extraction of text from various document formats:
- PDF files (with OCR fallback for scanned documents)
- DOCX files (with table and image support)
- DOC files (legacy format)
- RTF files
- ODT files (OpenDocument)

Features:
- Async file I/O for non-blocking operations
- Concurrent extraction for batch processing
- Progress callbacks for large documents
- Memory-efficient streaming for large files
- OCR fallback for scanned PDFs
"""

import asyncio
import hashlib
import logging
from concurrent.futures import ThreadPoolExecutor
from dataclasses import dataclass, field
from enum import Enum
from pathlib import Path
from typing import AsyncIterator, Callable, Dict, List, Optional, Any
import io

logger = logging.getLogger(__name__)

# Optional dependencies
try:
    from PyPDF2 import PdfReader
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False
    logger.warning("PyPDF2 not installed. PDF extraction disabled.")

try:
    import docx
    from docx.table import Table
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False
    logger.warning("python-docx not installed. DOCX extraction disabled.")

try:
    import pytesseract
    from PIL import Image
    OCR_SUPPORT = True
except ImportError:
    OCR_SUPPORT = False
    logger.warning("pytesseract not installed. OCR fallback disabled.")

try:
    import fitz  # PyMuPDF for better PDF handling
    PYMUPDF_SUPPORT = True
except ImportError:
    PYMUPDF_SUPPORT = False

try:
    from striprtf.striprtf import rtf_to_text
    RTF_SUPPORT = True
except ImportError:
    RTF_SUPPORT = False

try:
    from odf import text as odf_text
    from odf.opendocument import load as load_odt
    ODT_SUPPORT = True
except ImportError:
    ODT_SUPPORT = False


class ExtractionStatus(Enum):
    """Status of document extraction."""
    PENDING = "pending"
    IN_PROGRESS = "in_progress"
    COMPLETED = "completed"
    FAILED = "failed"
    PARTIAL = "partial"  # Some pages failed


@dataclass
class ExtractionResult:
    """Result of document extraction."""
    file_path: str
    status: ExtractionStatus
    text: str = ""
    page_count: int = 0
    extracted_pages: int = 0
    word_count: int = 0
    char_count: int = 0
    content_hash: str = ""
    metadata: Dict[str, Any] = field(default_factory=dict)
    error: Optional[str] = None
    extraction_time_ms: float = 0.0

    def __post_init__(self):
        if self.text:
            self.word_count = len(self.text.split())
            self.char_count = len(self.text)
            self.content_hash = hashlib.sha256(self.text.encode()).hexdigest()


@dataclass
class PageResult:
    """Result from extracting a single page."""
    page_number: int
    text: str
    success: bool
    error: Optional[str] = None


ProgressCallback = Callable[[int, int, str], None]  # (current, total, message)


class DocumentExtractor:
    """
    Async document extraction with support for multiple formats.
    Uses a thread pool for CPU-bound extraction operations.
    """

    SUPPORTED_EXTENSIONS = {
        '.pdf': 'pdf',
        '.docx': 'docx',
        '.doc': 'doc',
        '.rtf': 'rtf',
        '.odt': 'odt',
    }

    def __init__(
        self,
        max_workers: int = 4,
        enable_ocr: bool = True,
        ocr_languages: str = "eng",
        max_file_size: int = 100 * 1024 * 1024,  # 100MB
    ):
        """
        Initialize document extractor.

        Args:
            max_workers: Maximum threads for extraction
            enable_ocr: Enable OCR fallback for scanned PDFs
            ocr_languages: Tesseract language codes
            max_file_size: Maximum file size to process
        """
        self.max_workers = max_workers
        self.enable_ocr = enable_ocr and OCR_SUPPORT
        self.ocr_languages = ocr_languages
        self.max_file_size = max_file_size
        self._executor = ThreadPoolExecutor(max_workers=max_workers)

    def __del__(self):
        """Cleanup executor on destruction."""
        self._executor.shutdown(wait=False)

    def get_supported_formats(self) -> Dict[str, bool]:
        """Get dictionary of supported formats and their availability."""
        return {
            'pdf': PDF_SUPPORT or PYMUPDF_SUPPORT,
            'docx': DOCX_SUPPORT,
            'doc': False,  # Requires external tools
            'rtf': RTF_SUPPORT,
            'odt': ODT_SUPPORT,
            'ocr': OCR_SUPPORT,
        }

    def is_supported(self, file_path: Path) -> bool:
        """Check if file type is supported."""
        ext = file_path.suffix.lower()
        if ext not in self.SUPPORTED_EXTENSIONS:
            return False

        format_type = self.SUPPORTED_EXTENSIONS[ext]
        support_map = self.get_supported_formats()
        return support_map.get(format_type, False)

    async def extract(
        self,
        file_path: Path,
        progress_callback: Optional[ProgressCallback] = None
    ) -> ExtractionResult:
        """
        Extract text from a document asynchronously.

        Args:
            file_path: Path to document
            progress_callback: Optional callback for progress updates

        Returns:
            ExtractionResult with extracted text and metadata
        """
        import time
        start_time = time.time()

        file_path = Path(file_path)

        # Validate file
        if not file_path.exists():
            return ExtractionResult(
                file_path=str(file_path),
                status=ExtractionStatus.FAILED,
                error="File not found"
            )

        if not file_path.is_file():
            return ExtractionResult(
                file_path=str(file_path),
                status=ExtractionStatus.FAILED,
                error="Path is not a file"
            )

        try:
            file_size = file_path.stat().st_size
            if file_size > self.max_file_size:
                return ExtractionResult(
                    file_path=str(file_path),
                    status=ExtractionStatus.FAILED,
                    error=f"File too large: {file_size} bytes (max: {self.max_file_size})"
                )
        except OSError as e:
            return ExtractionResult(
                file_path=str(file_path),
                status=ExtractionStatus.FAILED,
                error=f"Cannot access file: {e}"
            )

        ext = file_path.suffix.lower()

        if ext not in self.SUPPORTED_EXTENSIONS:
            return ExtractionResult(
                file_path=str(file_path),
                status=ExtractionStatus.FAILED,
                error=f"Unsupported file type: {ext}"
            )

        # Run extraction in thread pool
        loop = asyncio.get_event_loop()

        try:
            if ext == '.pdf':
                result = await loop.run_in_executor(
                    self._executor,
                    lambda: self._extract_pdf_sync(file_path, progress_callback)
                )
            elif ext == '.docx':
                result = await loop.run_in_executor(
                    self._executor,
                    lambda: self._extract_docx_sync(file_path, progress_callback)
                )
            elif ext == '.rtf':
                result = await loop.run_in_executor(
                    self._executor,
                    lambda: self._extract_rtf_sync(file_path)
                )
            elif ext == '.odt':
                result = await loop.run_in_executor(
                    self._executor,
                    lambda: self._extract_odt_sync(file_path)
                )
            else:
                result = ExtractionResult(
                    file_path=str(file_path),
                    status=ExtractionStatus.FAILED,
                    error=f"No extractor for: {ext}"
                )

            result.extraction_time_ms = (time.time() - start_time) * 1000
            return result

        except Exception as e:
            logger.exception(f"Extraction failed for {file_path}")
            return ExtractionResult(
                file_path=str(file_path),
                status=ExtractionStatus.FAILED,
                error=str(e),
                extraction_time_ms=(time.time() - start_time) * 1000
            )

    def _extract_pdf_sync(
        self,
        file_path: Path,
        progress_callback: Optional[ProgressCallback] = None
    ) -> ExtractionResult:
        """Synchronous PDF extraction."""
        if PYMUPDF_SUPPORT:
            return self._extract_pdf_pymupdf(file_path, progress_callback)
        elif PDF_SUPPORT:
            return self._extract_pdf_pypdf2(file_path, progress_callback)
        else:
            return ExtractionResult(
                file_path=str(file_path),
                status=ExtractionStatus.FAILED,
                error="No PDF library available"
            )

    def _extract_pdf_pymupdf(
        self,
        file_path: Path,
        progress_callback: Optional[ProgressCallback] = None
    ) -> ExtractionResult:
        """Extract PDF using PyMuPDF (higher quality)."""
        try:
            doc = fitz.open(file_path)
            page_count = len(doc)
            text_parts = []
            extracted_pages = 0
            metadata = {}

            # Get document metadata
            if doc.metadata:
                metadata = {
                    'title': doc.metadata.get('title', ''),
                    'author': doc.metadata.get('author', ''),
                    'subject': doc.metadata.get('subject', ''),
                    'creator': doc.metadata.get('creator', ''),
                }

            for i, page in enumerate(doc):
                if progress_callback:
                    progress_callback(i + 1, page_count, f"Extracting page {i + 1}")

                text = page.get_text()

                # If no text and OCR is enabled, try OCR
                if not text.strip() and self.enable_ocr:
                    try:
                        # Render page to image
                        pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))  # 2x zoom for better OCR
                        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                        text = pytesseract.image_to_string(img, lang=self.ocr_languages)
                    except Exception as e:
                        logger.warning(f"OCR failed for page {i + 1}: {e}")

                if text.strip():
                    text_parts.append(text)
                    extracted_pages += 1

            doc.close()

            full_text = "\n\n".join(text_parts)
            status = ExtractionStatus.COMPLETED if extracted_pages == page_count else ExtractionStatus.PARTIAL

            return ExtractionResult(
                file_path=str(file_path),
                status=status,
                text=full_text,
                page_count=page_count,
                extracted_pages=extracted_pages,
                metadata=metadata
            )

        except Exception as e:
            return ExtractionResult(
                file_path=str(file_path),
                status=ExtractionStatus.FAILED,
                error=str(e)
            )

    def _extract_pdf_pypdf2(
        self,
        file_path: Path,
        progress_callback: Optional[ProgressCallback] = None
    ) -> ExtractionResult:
        """Extract PDF using PyPDF2 (fallback)."""
        try:
            reader = PdfReader(file_path)
            page_count = len(reader.pages)
            text_parts = []
            extracted_pages = 0
            metadata = {}

            # Get document metadata
            if reader.metadata:
                metadata = {
                    'title': reader.metadata.get('/Title', ''),
                    'author': reader.metadata.get('/Author', ''),
                    'subject': reader.metadata.get('/Subject', ''),
                    'creator': reader.metadata.get('/Creator', ''),
                }

            for i, page in enumerate(reader.pages):
                if progress_callback:
                    progress_callback(i + 1, page_count, f"Extracting page {i + 1}")

                try:
                    text = page.extract_text()
                    if text and text.strip():
                        text_parts.append(text)
                        extracted_pages += 1
                except Exception as e:
                    logger.warning(f"Failed to extract page {i + 1}: {e}")

            full_text = "\n\n".join(text_parts)
            status = ExtractionStatus.COMPLETED if extracted_pages == page_count else ExtractionStatus.PARTIAL

            return ExtractionResult(
                file_path=str(file_path),
                status=status,
                text=full_text,
                page_count=page_count,
                extracted_pages=extracted_pages,
                metadata=metadata
            )

        except Exception as e:
            return ExtractionResult(
                file_path=str(file_path),
                status=ExtractionStatus.FAILED,
                error=str(e)
            )

    def _extract_docx_sync(
        self,
        file_path: Path,
        progress_callback: Optional[ProgressCallback] = None
    ) -> ExtractionResult:
        """Synchronous DOCX extraction."""
        if not DOCX_SUPPORT:
            return ExtractionResult(
                file_path=str(file_path),
                status=ExtractionStatus.FAILED,
                error="python-docx not installed"
            )

        try:
            doc = docx.Document(file_path)
            text_parts = []
            metadata = {}

            # Get core properties
            if doc.core_properties:
                metadata = {
                    'title': doc.core_properties.title or '',
                    'author': doc.core_properties.author or '',
                    'subject': doc.core_properties.subject or '',
                    'created': str(doc.core_properties.created) if doc.core_properties.created else '',
                    'modified': str(doc.core_properties.modified) if doc.core_properties.modified else '',
                }

            # Extract paragraphs
            total_elements = len(doc.paragraphs) + len(doc.tables)
            current = 0

            for para in doc.paragraphs:
                current += 1
                if progress_callback:
                    progress_callback(current, total_elements, "Extracting paragraphs")

                if para.text.strip():
                    text_parts.append(para.text)

            # Extract tables
            for table in doc.tables:
                current += 1
                if progress_callback:
                    progress_callback(current, total_elements, "Extracting tables")

                table_text = self._extract_table(table)
                if table_text:
                    text_parts.append(table_text)

            full_text = "\n\n".join(text_parts)

            return ExtractionResult(
                file_path=str(file_path),
                status=ExtractionStatus.COMPLETED,
                text=full_text,
                page_count=1,  # DOCX doesn't have fixed pages
                extracted_pages=1,
                metadata=metadata
            )

        except Exception as e:
            return ExtractionResult(
                file_path=str(file_path),
                status=ExtractionStatus.FAILED,
                error=str(e)
            )

    def _extract_table(self, table: "Table") -> str:
        """Extract text from a DOCX table."""
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        return "\n".join(rows)

    def _extract_rtf_sync(self, file_path: Path) -> ExtractionResult:
        """Synchronous RTF extraction."""
        if not RTF_SUPPORT:
            return ExtractionResult(
                file_path=str(file_path),
                status=ExtractionStatus.FAILED,
                error="striprtf not installed"
            )

        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                rtf_content = f.read()

            text = rtf_to_text(rtf_content)

            return ExtractionResult(
                file_path=str(file_path),
                status=ExtractionStatus.COMPLETED,
                text=text,
                page_count=1,
                extracted_pages=1
            )

        except Exception as e:
            return ExtractionResult(
                file_path=str(file_path),
                status=ExtractionStatus.FAILED,
                error=str(e)
            )

    def _extract_odt_sync(self, file_path: Path) -> ExtractionResult:
        """Synchronous ODT extraction."""
        if not ODT_SUPPORT:
            return ExtractionResult(
                file_path=str(file_path),
                status=ExtractionStatus.FAILED,
                error="odfpy not installed"
            )

        try:
            doc = load_odt(file_path)
            text_parts = []

            for elem in doc.body.getElementsByType(odf_text.P):
                text = str(elem)
                if text.strip():
                    text_parts.append(text)

            full_text = "\n\n".join(text_parts)

            return ExtractionResult(
                file_path=str(file_path),
                status=ExtractionStatus.COMPLETED,
                text=full_text,
                page_count=1,
                extracted_pages=1
            )

        except Exception as e:
            return ExtractionResult(
                file_path=str(file_path),
                status=ExtractionStatus.FAILED,
                error=str(e)
            )

    async def extract_batch(
        self,
        file_paths: List[Path],
        max_concurrent: int = 4,
        progress_callback: Optional[Callable[[int, int, str], None]] = None
    ) -> List[ExtractionResult]:
        """
        Extract text from multiple documents concurrently.

        Args:
            file_paths: List of file paths
            max_concurrent: Maximum concurrent extractions
            progress_callback: Optional progress callback (file_index, total, filename)

        Returns:
            List of ExtractionResult objects
        """
        semaphore = asyncio.Semaphore(max_concurrent)
        results = []
        total = len(file_paths)

        async def extract_with_semaphore(i: int, path: Path) -> ExtractionResult:
            async with semaphore:
                if progress_callback:
                    progress_callback(i + 1, total, path.name)
                return await self.extract(path)

        tasks = [
            extract_with_semaphore(i, path)
            for i, path in enumerate(file_paths)
        ]

        results = await asyncio.gather(*tasks)
        return list(results)

    async def extract_pages(
        self,
        file_path: Path,
        page_numbers: Optional[List[int]] = None
    ) -> AsyncIterator[PageResult]:
        """
        Extract specific pages from a PDF as an async generator.

        Args:
            file_path: Path to PDF
            page_numbers: Specific page numbers to extract (1-indexed), or None for all

        Yields:
            PageResult for each page
        """
        file_path = Path(file_path)

        if file_path.suffix.lower() != '.pdf':
            yield PageResult(
                page_number=0,
                text="",
                success=False,
                error="Only PDF files support page-by-page extraction"
            )
            return

        if not PDF_SUPPORT and not PYMUPDF_SUPPORT:
            yield PageResult(
                page_number=0,
                text="",
                success=False,
                error="No PDF library available"
            )
            return

        loop = asyncio.get_event_loop()

        def extract_page(page_num: int) -> PageResult:
            try:
                if PYMUPDF_SUPPORT:
                    doc = fitz.open(file_path)
                    if page_num < 1 or page_num > len(doc):
                        doc.close()
                        return PageResult(
                            page_number=page_num,
                            text="",
                            success=False,
                            error=f"Page {page_num} out of range"
                        )
                    page = doc[page_num - 1]
                    text = page.get_text()
                    doc.close()
                else:
                    reader = PdfReader(file_path)
                    if page_num < 1 or page_num > len(reader.pages):
                        return PageResult(
                            page_number=page_num,
                            text="",
                            success=False,
                            error=f"Page {page_num} out of range"
                        )
                    page = reader.pages[page_num - 1]
                    text = page.extract_text() or ""

                return PageResult(
                    page_number=page_num,
                    text=text,
                    success=True
                )

            except Exception as e:
                return PageResult(
                    page_number=page_num,
                    text="",
                    success=False,
                    error=str(e)
                )

        # Determine pages to extract
        if page_numbers is None:
            # Get page count
            if PYMUPDF_SUPPORT:
                doc = fitz.open(file_path)
                page_count = len(doc)
                doc.close()
            else:
                reader = PdfReader(file_path)
                page_count = len(reader.pages)
            page_numbers = list(range(1, page_count + 1))

        for page_num in page_numbers:
            result = await loop.run_in_executor(
                self._executor,
                lambda pn=page_num: extract_page(pn)
            )
            yield result


# Module-level convenience functions

_default_extractor: Optional[DocumentExtractor] = None


def get_extractor() -> DocumentExtractor:
    """Get the default document extractor instance."""
    global _default_extractor
    if _default_extractor is None:
        _default_extractor = DocumentExtractor()
    return _default_extractor


async def extract_document(file_path: Path) -> ExtractionResult:
    """Extract text from a document using the default extractor."""
    return await get_extractor().extract(file_path)


async def extract_documents(file_paths: List[Path]) -> List[ExtractionResult]:
    """Extract text from multiple documents using the default extractor."""
    return await get_extractor().extract_batch(file_paths)


# CLI interface
if __name__ == "__main__":
    import argparse
    import sys

    parser = argparse.ArgumentParser(description="Async Document Extractor")
    parser.add_argument("files", nargs="+", help="Files to extract")
    parser.add_argument("--output", "-o", help="Output file (default: stdout)")
    parser.add_argument("--workers", type=int, default=4, help="Max worker threads")
    parser.add_argument("--no-ocr", action="store_true", help="Disable OCR fallback")
    parser.add_argument("--formats", action="store_true", help="List supported formats")
    args = parser.parse_args()

    if args.formats:
        extractor = DocumentExtractor()
        print("\nSupported formats:")
        for fmt, available in extractor.get_supported_formats().items():
            status = "✓" if available else "✗"
            print(f"  {status} {fmt}")
        sys.exit(0)

    async def main():
        extractor = DocumentExtractor(
            max_workers=args.workers,
            enable_ocr=not args.no_ocr
        )

        file_paths = [Path(f) for f in args.files]

        def progress(current, total, msg):
            print(f"[{current}/{total}] {msg}", file=sys.stderr)

        results = await extractor.extract_batch(file_paths, progress_callback=progress)

        output = []
        for result in results:
            output.append(f"\n{'='*60}")
            output.append(f"File: {result.file_path}")
            output.append(f"Status: {result.status.value}")
            output.append(f"Pages: {result.extracted_pages}/{result.page_count}")
            output.append(f"Words: {result.word_count}")
            output.append(f"Time: {result.extraction_time_ms:.2f}ms")
            if result.error:
                output.append(f"Error: {result.error}")
            output.append(f"{'='*60}")
            if result.text:
                output.append(result.text[:2000] + "..." if len(result.text) > 2000 else result.text)

        text = "\n".join(output)

        if args.output:
            with open(args.output, 'w', encoding='utf-8') as f:
                f.write(text)
            print(f"Output written to {args.output}", file=sys.stderr)
        else:
            print(text)

    asyncio.run(main())
